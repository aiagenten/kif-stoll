"""Generate KIF STOLL contract as .docx based on AI Agenten template v2."""
import copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

TEMPLATE = "/Users/buddymoltbot/clawd/ideas/Oppdragsavtale-AI-Agenten-mal-v2.docx"
OUTPUT = "/Users/buddymoltbot/clawd/repos/aiagenten/kif-stoll/kontrakt/Oppdragsavtale-KIF-STOLL.docx"

# Load template to get styles, images, headers/footers
template = Document(TEMPLATE)

# Clone the template (preserves all styles, headers, footers, images)
doc = Document(TEMPLATE)

# Clear all existing content (paragraphs + tables)
for p in doc.paragraphs:
    p._element.getparent().remove(p._element)
for t in doc.tables:
    t._element.getparent().remove(t._element)

# Helper functions
def add_para(text, style='normal', bold=False, alignment=None, color=None, size=None, space_after=None):
    p = doc.add_paragraph(text, style=style)
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    for run in p.runs:
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = RGBColor(*color)
        if size:
            run.font.size = Pt(size)
    return p

def add_heading(text, level=2):
    return doc.add_heading(text, level=level)

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_table_simple(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
    return table

# ==================== BUILD CONTRACT ====================

# Cover page
add_para('', space_after=60)
add_para('OPPDRAGSAVTALE', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=28, space_after=12)
add_para('AI Agenten AS', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, space_after=6)
add_para('', space_after=12)
add_para('Webutvikling · AI-automatisering · Digitale tjenester', alignment=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=40)
add_para('', space_after=40)
add_para('Versjon 1.0 — Mars 2026', alignment=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=6)

doc.add_page_break()

# === Main contract ===
add_para('OPPDRAGSAVTALE', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18, space_after=12)
add_para('Inngått mellom:', space_after=12)

# Parties table
parties = doc.add_table(rows=2, cols=2)
parties.alignment = WD_TABLE_ALIGNMENT.CENTER
parties.rows[0].cells[0].text = 'LEVERANDØR'
parties.rows[0].cells[1].text = 'KUNDE'
for cell in parties.rows[0].cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

parties.rows[1].cells[0].text = (
    'AI Agenten AS\n'
    'Org.nr.: 933 548 705\n'
    'kontakt@aiagenten.no\n'
    'https://aiagenten.no'
)
parties.rows[1].cells[1].text = (
    'Kongsberg Idrettsforening v/ STOLL Esport\n'
    'Org.nr.: 984 060 745\n'
    'hallo@stoll.gg\n'
    'https://stoll.gg'
)

add_para('')
add_para('Heretter samlet omtalt som «partene».', space_after=12)

# § 1
add_heading('§ 1. AVTALENS FORMÅL OG OMFANG')
add_para('1.1  Denne avtalen regulerer leverandørens levering av digitale tjenester til kunden som beskrevet i Vedlegg A — Oppdragsbeskrivelse.')
add_para('1.2  Tjenestene omfatter utvikling av ny nettside med CMS-system, AI-funksjoner, eventkalender med booking, og løpende serviceavtale som nærmere beskrevet i Vedlegg A.')
add_para('1.3  Alt arbeid utover det som er spesifikt beskrevet i Vedlegg A regnes som tilleggsarbeid og prises separat i henhold til § 6.')

# § 2
add_heading('§ 2. LEVERANSER OG GODKJENNING')
add_para('2.1  Leverandøren skal levere de produkter og tjenester som er angitt i Vedlegg A innen avtalte frister.')
add_para('2.2  Kunden skal gjennomgå og godkjenne leveranser innen 10 arbeidsdager etter mottak. Dersom tilbakemelding ikke er gitt innen fristen, anses leveransen som godkjent.')
add_para('2.3  Leverandøren kan ikke holdes ansvarlig for forsinkelser som skyldes manglende tilbakemeldinger, godkjenninger eller materiale fra kunden.')
add_para('2.4  Mindre feil som ikke påvirker funksjonaliteten, gir ikke kunden rett til å nekte godkjenning, men skal utbedres av leverandøren innen rimelig tid.')

# § 3
add_heading('§ 3. KUNDENS FORPLIKTELSER')
add_para('3.1  Kunden forplikter seg til å:')
add_bullet('Stille nødvendig informasjon, materiale og tilganger til rådighet i tide')
add_bullet('Utpeke en kontaktperson med beslutningsmyndighet')
add_bullet('Delta i avtalte møter og gjennomganger')
add_bullet('Gi skriftlig (e-post holder) tilbakemelding innen avtalte frister')
add_para('3.2  Forsinkelser, ekstraarbeid eller merkostnader som følge av kundens manglende oppfyllelse av § 3.1 belastes kunden som tilleggsarbeid etter timepris.')

# § 4
add_heading('§ 4. TIDSPLAN OG MILEPÆLER')
add_para('4.1  Prosjektets estimerte tidsplan fremgår av Vedlegg A. Tidsplanen er basert på at kunden leverer nødvendig materiale og godkjenninger til avtalt tid.')
add_para('4.2  Dersom kunden ber om endringer underveis som påvirker tidsplanen, justeres frister tilsvarende.')
add_para('4.3  Leverandøren har rett til å avbryte prosjektet og kreve betaling for utført arbeid dersom kunden ikke svarer på henvendelser innen 15 arbeidsdager.')

# § 5
add_heading('§ 5. PRIS OG BETALINGSVILKÅR')
add_para('5.1  Prosjektets totale pris fremgår av Vedlegg B — Prisoversikt. Alle priser er oppgitt eks. MVA.')
add_para('5.2  Betaling skjer etter følgende plan:')
add_bullet('50 % ved signering av avtalen (kr 9 500,-)')
add_bullet('50 % ved endelig leveranse og godkjenning (kr 9 500,-)')
add_para('5.3  Løpende serviceavtale faktureres årlig forskudd med oppstart fra leveransedato.')
add_para('5.4  Betalingsfrist er 14 dager fra fakturadato.')
add_para('5.5  Ved forsinket betaling beregnes forsinkelsesrente i henhold til forsinkelsesrenteloven.')
add_para('5.6  Leverandøren kan stanse videre arbeid og holde tilbake leveranser ved mislighold av betalingsforpliktelser.')

# § 6
add_heading('§ 6. ENDRINGER OG TILLEGGSARBEID')
add_para('6.1  Endringer i avtalens omfang skal avtales skriftlig mellom partene før arbeidet påbegynnes.')
add_para('6.2  Tilleggsarbeid utover Vedlegg A prises etter leverandørens gjeldende timepris, som fremgår av Vedlegg B.')
add_para('6.3  Leverandøren forbeholder seg retten til å avvise endringsønsker dersom disse er uforholdsmessig ressurskrevende eller teknisk uforsvarlige.')

# § 7
add_heading('§ 7. IMMATERIELLE RETTIGHETER (IP)')
add_para('7.1  Kundens rettigheter: Etter full betaling overdras eiendomsretten til det ferdige, tilpassede materialet (design, tekst, tilpasset kode) som er spesifikt utviklet for kunden.')
add_para('7.2  Leverandørens bakgrunns-IP: Leverandøren beholder alle rettigheter til egenutviklede verktøy, metoder, rammeverk, biblioteker og generiske løsninger som brukes i prosjektet. Kunden får en ikke-eksklusiv bruksrett til disse.')
add_para('7.3  Tredjepartskomponenter (open source, lisensierte verktøy) er underlagt egne lisensvilkår.')
add_para('7.4  Leverandøren har rett til å vise til prosjektet som referanse og bruke det i markedsføring, med mindre kunden skriftlig ber om konfidensialitet.')

# § 8
add_heading('§ 8. KONFIDENSIALITET')
add_para('8.1  Begge parter forplikter seg til å behandle fortrolig informasjon konfidensielt.')
add_para('8.2  Konfidensialitetsplikten gjelder i 2 år etter avtalens opphør.')
add_para('8.3  Taushetsplikten gjelder ikke informasjon som er offentlig tilgjengelig, eller som parten kan vise å ha mottatt fra en tredjepart uten konfidensialitetsforpliktelse.')

# § 9
add_heading('§ 9. PERSONVERN OG GDPR')
add_para('9.1  Partene forplikter seg til å overholde gjeldende personvernlovgivning (GDPR / personopplysningsloven).')
add_para('9.2  Dersom leverandøren behandler personopplysninger på vegne av kunden (kontaktskjema, booking-data, chatbot-meldinger), skal det inngås en separat databehandleravtale (Vedlegg D).')
add_para('9.3  Leverandørens behandling av personopplysninger reguleres av personvernerklæringen på https://aiagenten.no/personvern')

# § 10
add_heading('§ 10. ANSVARSBEGRENSNING')
add_para('10.1  Leverandørens samlede erstatningsansvar overfor kunden er begrenset til prosjektets totale kontraktsverdi ekskl. MVA.')
add_para('10.2  Leverandøren er ikke ansvarlig for:')
add_bullet('Indirekte tap, følgeskader eller tapte inntekter')
add_bullet('Tap som skyldes kundens egne handlinger eller unnlatelser')
add_bullet('Feil i tredjepartsløsninger (hosting, betalingsleverandører, API-er etc.)')
add_bullet('Nedetid hos tjenesteleverandører utenfor leverandørens kontroll')
add_bullet('Sikkerhetshendelser som skyldes at kunden ikke har fulgt leverandørens anbefalinger')
add_para('10.3  Kunden har plikt til å begrense sitt tap.')

# § 11
add_heading('§ 11. MISLIGHOLD OG HEVING')
add_para('11.1  En part kan heve avtalen ved vesentlig mislighold dersom misligholdet ikke er utbedret innen 14 dager etter skriftlig varsel.')
add_para('11.2  Vesentlig mislighold fra kundens side inkluderer: manglende betaling, ikke å stille nødvendige ressurser til rådighet, eller gjentatte forsinkelser.')
add_para('11.3  Vesentlig mislighold fra leverandørens side inkluderer: manglende levering uten saklig grunn eller grov uaktsomhet.')
add_para('11.4  Ved heving fra leverandørens side pga. kundens mislighold har leverandøren krav på betaling for alt utført arbeid frem til hevingstidspunktet.')

# § 12
add_heading('§ 12. OPPSIGELSE AV LØPENDE AVTALER')
add_para('12.1  Løpende tjenesteavtaler (vedlikehold, hosting-støtte, månedlige abonnementer) kan sies opp med 30 dagers skriftlig varsel av begge parter.')
add_para('12.2  Kunden har ikke krav på refusjon av forhåndsbetalt abonnement ved oppsigelse.')

# § 13
add_heading('§ 13. FORCE MAJEURE')
add_para('13.1  Ingen av partene er ansvarlige for manglende oppfyllelse av sine forpliktelser dersom dette skyldes forhold utenfor partenes kontroll (force majeure).')
add_para('13.2  Den rammede parten skal varsle den andre parten uten ugrunnet opphold.')

# § 14
add_heading('§ 14. VILKÅR OG PERSONVERN')
add_para('14.1  Denne avtalen suppleres av leverandørens generelle Vilkår for bruk og Personvernerklæring, tilgjengelig på https://aiagenten.no/vilkaar og https://aiagenten.no/personvern')
add_para('14.2  Ved motstrid mellom denne avtalen og de generelle vilkårene, går denne avtalen foran.')

# § 15
add_heading('§ 15. LOVVALG OG TVISTELØSNING')
add_para('15.1  Avtalen er underlagt norsk rett.')
add_para('15.2  Partene skal søke å løse tvister i minnelighet. Dersom enighet ikke oppnås, avgjøres tvisten ved Kongsberg tingrett som verneting.')

# SIGNATURER
doc.add_page_break()
add_heading('SIGNATURER', level=1)
add_para('Denne avtalen er inngått i to originaler, én til hver part.')
add_para('')
add_para('Sted og dato: _______________________')
add_para('')

sig_table = doc.add_table(rows=5, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sig_table.rows[0].cells[0].text = 'AI Agenten AS'
sig_table.rows[0].cells[1].text = 'Kongsberg Idrettsforening'
for cell in sig_table.rows[0].cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
sig_table.rows[1].cells[0].text = 'Underskrift: ______________________'
sig_table.rows[1].cells[1].text = 'Underskrift: ______________________'
sig_table.rows[2].cells[0].text = 'Navn: Andreas Storaas-Barsnes'
sig_table.rows[2].cells[1].text = 'Navn: ______________________'
sig_table.rows[3].cells[0].text = 'Stilling: Daglig leder'
sig_table.rows[3].cells[1].text = 'Stilling: ______________________'
sig_table.rows[4].cells[0].text = 'Dato: ______________________'
sig_table.rows[4].cells[1].text = 'Dato: ______________________'

# VEDLEGG A
doc.add_page_break()
add_heading('VEDLEGG A — OPPDRAGSBESKRIVELSE', level=1)
add_para('')
add_para('Prosjektnavn:', bold=True)
add_para('STOLL Esportsenter — Nettside og digitale tjenester')
add_para('')
add_para('Beskrivelse av leveranse:', bold=True)
add_para('')
add_para('A.1 Leveranse: Ny nettside med CMS', bold=True)
add_para('Utvikling av komplett, responsiv nettside for STOLL Esportsenter med følgende:')
add_para('')
add_para('Offentlig nettside:', bold=True)
add_bullet('Forside med hero, om oss, medlemskap-info, kontaktskjema')
add_bullet('«Hva skjer på STOLL» — eventkalender med kommende arrangementer')
add_bullet('Booking-funksjon for events og bursdagsfeiringer')
add_bullet('Blogg / nyheter')
add_bullet('Sponsoroversikt (rullerende logo-karusell)')
add_bullet('Optimalisert for mobil, nettbrett og desktop')
add_para('')
add_para('CMS Admin-panel (innlogging for STOLL-administrator):', bold=True)
add_bullet('Redigere tekst, bilder og videoer på alle sider')
add_bullet('Administrere blogginnlegg (opprett, rediger, publiser)')
add_bullet('Administrere events og kalender')
add_bullet('Se og håndtere bookinger (godkjenn/avvis)')
add_bullet('Administrere sponsorlogoer')
add_bullet('Se innkommende kontakthenvendelser')
add_bullet('SEO-innstillinger per side')
add_para('')
add_para('Teknisk plattform:', bold=True)
add_bullet('React / Vite frontend')
add_bullet('Supabase database og autentisering')
add_bullet('Hosting inkludert i serviceavtale')
add_bullet('SSL-sertifikat')

add_para('')
add_para('A.2 Leveranse: AI-funksjoner', bold=True)
add_para('')
add_para('KI Chatbot:', bold=True)
add_bullet('Integrert chat-widget på nettsiden')
add_bullet('Trenet på STOLL sin informasjon (medlemskap, priser, åpningstider, spill)')
add_bullet('Svarer automatisk på vanlige spørsmål 24/7')
add_bullet('Kvote: 500 chatmeldinger per måned')
add_para('')
add_para('KI Innholdsgenerering:', bold=True)
add_bullet('Generering av artikler/nyhetsinnlegg med AI')
add_bullet('Generering av bilder til artikler')
add_bullet('Kvote: 15 artikler og 20 bilder per måned')

add_para('')
add_para('A.3 Leveranse: Sosiale Medier-automatisering', bold=True)
add_para('')
add_para('SoMe-modul i admin-panel:', bold=True)
add_bullet('Koble Facebook- og Instagram-konto')
add_bullet('AI-genererte innleggsforslag basert på eventkalender')
add_bullet('Preview og godkjenning før publisering')
add_bullet('Automatiske event-påminnelser (3 dager og 1 dag før)')
add_bullet('Feedback-system som forbedrer innholdet over tid')

add_para('')
add_para('A.4 Hva er IKKE inkludert:', bold=True)
add_bullet('Domeneregistrering (kunden beholder sitt eksisterende domene)')
add_bullet('Innhold (tekst/bilder) utover det som genereres med AI-verktøyene')
add_bullet('Fysisk opplæring (digital gjennomgang inkludert)')
add_bullet('Integrasjoner med tredjeparter utover det som er beskrevet')
add_bullet('Migrasjon av innhold fra eksisterende WordPress-side')

add_para('')
add_para('A.5 Estimert tidsplan:', bold=True)
add_para('')
timeline = doc.add_table(rows=4, cols=2)
timeline.alignment = WD_TABLE_ALIGNMENT.CENTER
timeline.rows[0].cells[0].text = 'Milepæl'
timeline.rows[0].cells[1].text = 'Estimert'
for cell in timeline.rows[0].cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
timeline.rows[1].cells[0].text = 'Oppstart'
timeline.rows[1].cells[1].text = 'Ved signering'
timeline.rows[2].cells[0].text = 'Demoversjon'
timeline.rows[2].cells[1].text = '1–2 uker etter oppstart'
timeline.rows[3].cells[0].text = 'Endelig levering'
timeline.rows[3].cells[1].text = '3–4 uker etter oppstart'

# VEDLEGG B
doc.add_page_break()
add_heading('VEDLEGG B — PRISOVERSIKT', level=1)
add_para('')
add_para('Engangskostnader', bold=True)
add_para('')
price1 = doc.add_table(rows=4, cols=4)
price1.alignment = WD_TABLE_ALIGNMENT.CENTER
h = ['Post', 'Ordinær pris', 'Rabatt', 'Pris eks. MVA']
for i, t in enumerate(h):
    price1.rows[0].cells[i].text = t
    for p in price1.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True
rows_data = [
    ['Nettside med CMS-system', 'kr 18 000,-', '50% lag/forening', 'kr 9 000,-'],
    ['AI-etablering (chatbot + innhold + SoMe)', 'kr 20 000,-', '50% lag/forening', 'kr 10 000,-'],
    ['Sum engangskostnad', '', '', 'kr 19 000,-'],
]
for i, rd in enumerate(rows_data):
    for j, v in enumerate(rd):
        price1.rows[i+1].cells[j].text = v

add_para('')
add_para('Løpende kostnader (månedlig)', bold=True)
add_para('')
price2 = doc.add_table(rows=4, cols=3)
price2.alignment = WD_TABLE_ALIGNMENT.CENTER
h2 = ['Post', 'Ordinær pris', 'Pris eks. MVA']
for i, t in enumerate(h2):
    price2.rows[0].cells[i].text = t
    for p in price2.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True
rows2 = [
    ['Serviceavtale (hosting, backup, support, monitoring)', 'kr 399,-/mnd', 'kr 399,-/mnd'],
    ['AI-tjenester (chatbot + innholdsgenerering)', 'kr 1 499,-/mnd', 'kr 199,-/mnd'],
    ['Sum månedlig', '', 'kr 598,-/mnd'],
]
for i, rd in enumerate(rows2):
    for j, v in enumerate(rd):
        price2.rows[i+1].cells[j].text = v

add_para('')
add_para('Serviceavtale faktureres årlig forskudd: kr 7 176,- eks. MVA', size=10)
add_para('')
add_para('Tilleggsarbeid', bold=True)
add_para('')
price3 = doc.add_table(rows=3, cols=2)
price3.alignment = WD_TABLE_ALIGNMENT.CENTER
price3.rows[0].cells[0].text = 'Post'
price3.rows[0].cells[1].text = 'Pris eks. MVA'
for cell in price3.rows[0].cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
price3.rows[1].cells[0].text = 'Timepris utviklingsarbeid'
price3.rows[1].cells[1].text = 'kr 1 200,-/t'
price3.rows[2].cells[0].text = 'Timepris konsultering/rådgivning'
price3.rows[2].cells[1].text = 'kr 1 500,-/t'

add_para('')
add_para('Betalingsplan engangskostnad', bold=True)
add_para('')
bp = doc.add_table(rows=4, cols=2)
bp.alignment = WD_TABLE_ALIGNMENT.CENTER
bp.rows[0].cells[0].text = 'Tidspunkt'
bp.rows[0].cells[1].text = 'Beløp eks. MVA'
for cell in bp.rows[0].cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
bp.rows[1].cells[0].text = 'Ved signering'
bp.rows[1].cells[1].text = 'kr 9 500,-'
bp.rows[2].cells[0].text = 'Ved endelig leveranse'
bp.rows[2].cells[1].text = 'kr 9 500,-'
bp.rows[3].cells[0].text = 'Totalt'
bp.rows[3].cells[1].text = 'kr 19 000,-'

# VEDLEGG C
doc.add_page_break()
add_heading('VEDLEGG C — SERVICEAVTALE & SUPPORT', level=1)

add_heading('1. Om avtalen')
add_para('Med denne månedlige serviceavtalen slipper kunden å bekymre seg for det tekniske. AI Agenten sørger for at nettsiden alltid er oppdatert, sikker og fungerer som den skal.')

add_heading('2. Hva er inkludert')
add_para('Garantert oppetid og overvåking', bold=True)
add_para('Leverandøren overvåker og vedlikeholder nettsiden kontinuerlig for å sikre stabil og pålitelig drift. Ved nedetid eller tekniske feil, løses dette så raskt som mulig.')
add_para('')
add_para('Tilgang til redigering', bold=True)
add_para('Kunden skal til enhver tid ha fungerende tilgang til CMS-panelet for redigering av nettsidens innhold, inkludert bilder, tekst, events og blogginnlegg.')
add_para('')
add_para('Innholdsoppdateringer', bold=True)
add_para('Oppdatering av bilder, videoer, events og annet innhold etter behov — utført av leverandøren på forespørsel fra kunden.')
add_para('')
add_para('Teknisk support', bold=True)
add_para('Rask respons ved tekniske problemer. Kunden skal aldri stå alene med utfordringer knyttet til nettsiden. Henvendelser besvares normalt innen 24 timer på hverdager.')
add_para('')
add_para('SEO- og AEO-vedlikehold', bold=True)
add_para('Løpende optimalisering for søkemotorer og AI-basert søk, slik at kunden forblir synlig og relevant i søkeresultater.')

add_heading('3. Pris')
add_para('Serviceavtalen koster kr 399,– eks. MVA per måned (kr 4 788,- per år). AI-tjenester koster kr 199,- eks. MVA per måned. Total månedskostnad: kr 598,- eks. MVA. Beløpet faktureres årlig forskudd.')

add_heading('4. Varighet og oppsigelse')
add_para('Avtalen løper fra leveringsdato og fornyes automatisk månedlig. Avtalen kan sies opp av begge parter med 3 måneders skriftlig varsel.')

add_heading('5. Hva er IKKE inkludert')
add_bullet('Utvikling av ny funksjonalitet (prises som tilleggsarbeid iht. § 6 i hovedavtalen)')
add_bullet('Profesjonell fotografering eller videoproduksjon')
add_bullet('Innholdsproduksjon utover enkel oppdatering (f.eks. blogginnlegg, markedsføring)')
add_bullet('Tredjeparts hosting-kostnader (dekkes av leverandøren inntil videre)')

# VEDLEGG D
doc.add_page_break()
add_heading('VEDLEGG D — DATABEHANDLERAVTALE', level=1)

add_heading('D.1 Formål')
add_para('AI Agenten AS behandler personopplysninger på vegne av Kongsberg Idrettsforening i forbindelse med drift av nettside, kontaktskjema, booking-system og chatbot.')

add_heading('D.2 Kategorier av personopplysninger')
add_bullet('Navn, e-post, telefonnummer (kontaktskjema og booking)')
add_bullet('Chatmeldinger (chatbot)')
add_bullet('IP-adresse og besøksstatistikk (anonymisert via Umami)')

add_heading('D.3 Behandlingens varighet')
add_para('Personopplysninger behandles så lenge avtalen løper. Ved opphør slettes eller returneres data innen 30 dager etter kundens instruks.')

add_heading('D.4 Sikkerhetstiltak')
add_bullet('Kryptert dataoverføring (TLS)')
add_bullet('Tilgangskontroll med rollebasert autentisering')
add_bullet('Database hostet i EU (Frankfurt, eu-central-1)')
add_bullet('Regelmessig sikkerhetsgjennomgang')

add_heading('D.5 Underleverandører')
add_para('Leverandøren benytter følgende underleverandører for databehandling:')
add_para('')
sub_table = doc.add_table(rows=4, cols=3)
sub_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sub_table.rows[0].cells[0].text = 'Underleverandør'
sub_table.rows[0].cells[1].text = 'Formål'
sub_table.rows[0].cells[2].text = 'Lokasjon'
for cell in sub_table.rows[0].cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
sub_table.rows[1].cells[0].text = 'Supabase Inc.'
sub_table.rows[1].cells[1].text = 'Database og autentisering'
sub_table.rows[1].cells[2].text = 'EU (Frankfurt)'
sub_table.rows[2].cells[0].text = 'OpenAI Inc.'
sub_table.rows[2].cells[1].text = 'AI chatbot og innholdsgenerering'
sub_table.rows[2].cells[2].text = 'USA*'
sub_table.rows[3].cells[0].text = 'Netlify Inc.'
sub_table.rows[3].cells[1].text = 'Hosting'
sub_table.rows[3].cells[2].text = 'EU/Global CDN'

add_para('')
add_para('*OpenAI: Data brukes ikke til trening av modeller. Standard Contractual Clauses (SCC) gjelder for overføring til USA.', size=9)

add_heading('D.6 Rettigheter')
add_para('Kunden kan til enhver tid be om innsyn, retting eller sletting av personopplysninger som behandles på deres vegne.')

add_para('')
add_para('Kontrakten er utarbeidet med inspirasjon fra Statens standardavtaler (SSA-O) og Dataforeningens IT-kontraktstandarder, tilpasset SMB-markedet.', size=9)

# Save
doc.save(OUTPUT)
print(f"✅ Saved to {OUTPUT}")
